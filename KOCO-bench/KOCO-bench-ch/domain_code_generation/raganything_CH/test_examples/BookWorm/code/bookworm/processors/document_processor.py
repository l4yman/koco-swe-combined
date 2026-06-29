"""
Document processor for extracting text from various file formats
"""
import logging
from pathlib import Path
from typing import Union, List, Optional

# Optional imports for document processing
try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

from ..models import ProcessedDocument
from ..library import LibraryManager
from ..utils import BookWormConfig, is_supported_file


class DocumentProcessor:
    """Process documents and extract knowledge"""
    
    def __init__(self, config: BookWormConfig, library_manager: Optional[LibraryManager] = None):
        self.config = config
        self.working_dir = Path(config.working_dir)
        self.logger = logging.getLogger(__name__)
        self.library_manager = library_manager or LibraryManager(config)
    
    async def process_document(self, file_path: Union[str, Path]) -> Optional[ProcessedDocument]:
        """Process a single document and extract text"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            self.logger.error(f"File not found: {file_path}")
            return None
        
        if not is_supported_file(file_path):
            self.logger.warning(f"Unsupported file type: {file_path}")
            return None
        
        self.logger.info(f"Processing document: {file_path.name}")
        
        # Create processed document
        doc = ProcessedDocument(
            original_path=str(file_path.absolute()),
            file_type=file_path.suffix.lower(),
            file_size=file_path.stat().st_size,
            status="processing"
        )
        
        try:
            # Extract text based on file type
            doc.text_content = await self._extract_text(file_path)
            
            if doc.text_content:
                # Save processed text
                processed_dir = Path(self.config.processed_dir)
                processed_dir.mkdir(parents=True, exist_ok=True)
                processed_file = processed_dir / f"{doc.id}_{file_path.stem}.txt"
                processed_file.write_text(doc.text_content, encoding='utf-8')
                doc.processed_path = str(processed_file)
                doc.status = "completed"
                
                self.logger.info(f"✅ Successfully processed: {file_path.name}")
                return doc
            else:
                doc.status = "failed"
                doc.error_message = "No text content extracted"
                self.logger.error(f"❌ No text content extracted from: {file_path.name}")
                return None
                
        except Exception as e:
            doc.status = "failed"
            doc.error_message = str(e)
            self.logger.error(f"❌ Error processing {file_path.name}: {e}")
            return None
    
    async def _extract_text(self, file_path: Path) -> str:
        """Extract text from various file formats"""
        file_ext = file_path.suffix.lower()
        
        if file_ext == '.txt':
            return file_path.read_text(encoding='utf-8', errors='ignore')
        
        elif file_ext == '.md':
            return file_path.read_text(encoding='utf-8', errors='ignore')
        
        elif file_ext == '.pdf':
            return await self._extract_pdf_text(file_path)
        
        elif file_ext in ['.doc', '.docx']:
            return await self._extract_word_text(file_path)
        
        else:
            # Try to read as text
            try:
                return file_path.read_text(encoding='utf-8', errors='ignore')
            except Exception as e:
                self.logger.warning(f"Could not read {file_path} as text: {e}")
                return ""
    
    async def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF using multiple methods"""
        text_content = ""
        
        # Try PyMuPDF first
        if fitz:
            try:
                doc = fitz.open(str(file_path))
                for page_num in range(doc.page_count):
                    page = doc.load_page(page_num)
                    # Use explicit text mode to satisfy type checkers/newer fitz APIs
                    text_content += page.get_text("text") + "\n\n"  # type: ignore[attr-defined]
                doc.close()
                
                if text_content.strip():
                    self.logger.info(f"Extracted text using PyMuPDF: {len(text_content)} chars")
                    return text_content
            except Exception as e:
                self.logger.warning(f"PyMuPDF extraction failed: {e}")
        
        # Fallback to pdfplumber
        if pdfplumber:
            try:
                with pdfplumber.open(str(file_path)) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_content += page_text + "\n\n"
                
                if text_content.strip():
                    self.logger.info(f"Extracted text using pdfplumber: {len(text_content)} chars")
                    return text_content
            except Exception as e:
                self.logger.warning(f"pdfplumber extraction failed: {e}")
        
        self.logger.error("All PDF extraction methods failed")
        return ""
    
    async def _extract_word_text(self, file_path: Path) -> str:
        """Extract text from Word documents"""
        if not DocxDocument:
            self.logger.error("python-docx not available for Word document processing")
            return ""
        
        try:
            doc = DocxDocument(str(file_path))
            text_content = ""
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            self.logger.info(f"Extracted text from Word document: {len(text_content)} chars")
            return text_content
            
        except Exception as e:
            self.logger.error(f"Word document extraction failed: {e}")
            return ""
    
    async def process_directory(self, directory_path: Union[str, Path]) -> List[ProcessedDocument]:
        """Process all supported documents in a directory"""
        directory_path = Path(directory_path)
        documents = []
        
        if not directory_path.exists():
            self.logger.error(f"Directory not found: {directory_path}")
            return documents
        
        self.logger.info(f"Processing documents in: {directory_path}")
        
        for file_path in directory_path.rglob("*"):
            if file_path.is_file() and is_supported_file(file_path):
                processed_doc = await self.process_document(file_path)
                if processed_doc:
                    documents.append(processed_doc)
        
        self.logger.info(f"Processed {len(documents)} documents from {directory_path}")
        return documents
    
    async def process_directory_as_single_document(self, directory_path: Union[str, Path]) -> Optional[ProcessedDocument]:
        """Process an entire directory as a single document (for Obsidian vaults, doc collections)"""
        directory_path = Path(directory_path)
        
        if not directory_path.exists():
            self.logger.error(f"Directory not found: {directory_path}")
            return None
        
        if not directory_path.is_dir():
            self.logger.error(f"Path is not a directory: {directory_path}")
            return None
        
        self.logger.info(f"Processing directory as single document: {directory_path.name}")
        
        # Create processed document for the directory
        doc = ProcessedDocument(
            original_path=str(directory_path.absolute()),
            file_type="directory",
            file_size=0,
            status="processing"
        )
        
        try:
            # Collect all supported files in the directory
            file_paths = []
            total_size = 0
            combined_text = []
            
            for file_path in directory_path.rglob("*"):
                if file_path.is_file() and is_supported_file(file_path):
                    file_paths.append(str(file_path.relative_to(directory_path)))
                    total_size += file_path.stat().st_size
                    
                    # Extract text from each file
                    try:
                        file_text = await self._extract_text(file_path)
                        if file_text:
                            # Add file header to maintain context
                            relative_path = file_path.relative_to(directory_path)
                            combined_text.append(f"\n# File: {relative_path}\n\n{file_text}\n")
                            self.logger.debug(f"Added text from: {relative_path}")
                    except Exception as e:
                        self.logger.warning(f"Failed to extract text from {file_path}: {e}")
                        continue
            
            if not combined_text:
                doc.status = "failed"
                doc.error_message = "No extractable text found in directory"
                self.logger.error(f"❌ No extractable text found in: {directory_path.name}")
                return None
            
            # Combine all text content
            doc.text_content = "\n".join(combined_text)
            doc.file_size = total_size
            
            # Store metadata about the directory
            doc.metadata = {
                "is_directory": True,
                "sub_files": file_paths,
                "file_count": len(file_paths),
                "directory_name": directory_path.name,
                "processing_type": "directory_collection"
            }
            
            # Save processed text
            processed_dir = Path(self.config.processed_dir)
            processed_dir.mkdir(parents=True, exist_ok=True)
            processed_file = processed_dir / f"{doc.id}_{directory_path.name}.txt"
            processed_file.write_text(doc.text_content, encoding='utf-8')
            doc.processed_path = str(processed_file)
            doc.status = "completed"
            
            self.logger.info(f"✅ Successfully processed directory: {directory_path.name} ({len(file_paths)} files)")
            return doc
            
        except Exception as e:
            doc.status = "failed"
            doc.error_message = str(e)
            self.logger.error(f"❌ Error processing directory {directory_path.name}: {e}")
            return None
